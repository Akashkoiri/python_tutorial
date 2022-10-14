import docx

doc = docx.Document()

## Don't overwrite a pre_written file or you will loose all content 
## (solution:save it as a new file)

doc.add_paragraph().add_run('This is text')              

obj = doc.add_paragraph()

obj.add_run('Hello world!')

doc.add_picture('C:\\Users\\AKASH-PC\\Pictures\\Camera Roll\\img.jpg',height=docx.shared.Inches(6),width=docx.shared.Inches(5))

doc.add_picture('img\\img.jpg',height=docx.shared.Inches(6),width=docx.shared.Inches(5))


###  Making a Table:

records = (
    (3,'101','Spam'),
    (7,'422','Eggs'),
    (4,'631','Spam,spam,eggs,and spam')
)
table = doc.add_table(rows=2,cols=3)

hdr_cells = table.rows[0].cells

hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

for qty,id,desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[0].text = id
    row_cells[0].text = desc


doc.save('test.docx')
