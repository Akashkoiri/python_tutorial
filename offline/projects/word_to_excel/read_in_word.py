import docx

doc = docx.Document("test.docx")

#print(doc.paragraphs[1].text)
#print(len(doc.paragraphs))


#for i in range(len(doc.paragraphs)):
#    print(doc.paragraphs[i].text)

#--Concept of runs (diffrent styling)--#

#print(doc.paragraphs[0].runs)

#print(doc.paragraphs[0].runs[2].text)

#--Printing table--#

#  for table in doc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             print(cell.text)